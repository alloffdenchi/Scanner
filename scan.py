import cv2
import numpy as np
import imutils
from PIL import Image, ImageTk
from tkinter.ttk import Label
import tkinter as tk
import os
from google.cloud import vision
import io
from google.oauth2 import service_account
from tkinter import filedialog
from tkinter import messagebox
import threading
from tkinter import ttk
import webbrowser
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

def convert_cv_to_tkinter(cv_image):
    if len(cv_image.shape) == 3:
        image = cv2.cvtColor(cv_image, cv2.COLOR_BGR2RGB)
    else:
        image = cv2.cvtColor(cv_image, cv2.COLOR_GRAY2RGB)
    pil_image = Image.fromarray(image)
    return ImageTk.PhotoImage(pil_image)
# show image
def show_image(image, title, window):
    top = tk.Toplevel(window)
    top.title(title)
    photo = convert_cv_to_tkinter(image)
    label = tk.Label(top, image=photo)
    label.image = photo
    label.pack()
# preprocess image
def preprocess_image(image):
        # Resize image to fit screen
        screen_width = 1366
        screen_height = 768
        h, w = image.shape[:2]
        scale = min(screen_width / w, screen_height / h)
        new_w = int(w * scale)
        new_h = int(h * scale)
        return cv2.resize(image, (new_w, new_h))

def enhance_image_quality(image):
    """Cải thiện chất lượng ảnh trước khi xử lý"""
    # Chuyển sang ảnh xám
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Cân bằng histogram thích ứng để cải thiện độ tương phản
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
    contrast = clahe.apply(gray)
    
    # Giảm nhiễu bằng Non-local Means Denoising
    denoised = cv2.fastNlMeansDenoising(contrast, 
                                       h=10,  # Filtering strength
                                       templateWindowSize=7,
                                       searchWindowSize=21)
    
    # Tăng độ sắc nét
    kernel = np.array([[-1,-1,-1],
                      [-1, 9,-1],
                      [-1,-1,-1]])
    sharpened = cv2.filter2D(denoised, -1, kernel)
    
    return sharpened

def find_document_corners(image):  
    """Cải thiện phát hiện biên tài liệu"""
    # Tiền xử lý ảnh
    processed = enhance_image_quality(image)
    
    # Phát hiện cạnh với Canny cải tiến
    edge = cv2.Canny(processed, 30, 200)
    
    # Tăng kích thước kernel và số lần lặp để kết nối các đường đứt
    kernel = np.ones((5,5), np.uint8)
    edge = cv2.dilate(edge, kernel, iterations=2)
    edge = cv2.erode(edge, kernel, iterations=1)

    # Tìm contours
    cnts = cv2.findContours(edge, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    cnts = imutils.grab_contours(cnts)

    if not cnts:
        return None
    
    # Sắp xếp contours theo diện tích giảm dần
    cnts = sorted(cnts, key=cv2.contourArea, reverse=True)
    
    # Giảm ngưỡng diện tích tối thiểu xuống 3%
    min_area = image.shape[0] * image.shape[1] * 0.03
    
    for c in cnts:
        if cv2.contourArea(c) < min_area:
            continue
            
        # Tìm đa giác xấp xỉ với epsilon thích ứng
        peri = cv2.arcLength(c, True)
        epsilon = 0.02 * peri
        approx = cv2.approxPolyDP(c, epsilon, True)
        
        # Nếu tìm được 4 điểm và tỷ lệ cạnh hợp lý
        if len(approx) == 4 and verify_aspect_ratio(approx):
            return approx
            
        # Thử với nhiều giá trị epsilon khác nhau
        epsilon_values = [0.01, 0.02, 0.03, 0.04, 0.05]
        for eps in epsilon_values:
            approx = cv2.approxPolyDP(c, eps * peri, True)
            if len(approx) == 4 and verify_aspect_ratio(approx):
                return approx
    
    return None

def verify_aspect_ratio(corners):
    """Kiểm tra tỷ lệ cạnh của tài liệu có hợp lý không"""
    rect = order_points(corners.reshape(4, 2))
    (tl, tr, br, bl) = rect
    
    # Tính chiều dài các cạnh
    width_top = np.sqrt(((tr[0] - tl[0]) ** 2) + ((tr[1] - tl[1]) ** 2))
    width_bottom = np.sqrt(((br[0] - bl[0]) ** 2) + ((br[1] - bl[1]) ** 2))
    height_left = np.sqrt(((bl[0] - tl[0]) ** 2) + ((bl[1] - tl[1]) ** 2))
    height_right = np.sqrt(((br[0] - tr[0]) ** 2) + ((br[1] - tr[1]) ** 2))
    
    # Tỷ lệ chiều dài/rộng hợp lý cho tài liệu
    min_ratio = 0.5  # Cho phép tài liệu ngang
    max_ratio = 2.0  # Cho phép tài liệu dọc
    
    width = max(width_top, width_bottom)
    height = max(height_left, height_right)
    ratio = max(width, height) / min(width, height)
    
    return min_ratio <= ratio <= max_ratio

def order_points(pts):
    rect = np.zeros((4, 2), dtype="float32")
    s = pts.sum(axis=1)
    rect[0] = pts[np.argmin(s)]  # Top-left
    rect[2] = pts[np.argmax(s)]  # Bottom-right
    
    diff = np.diff(pts, axis=1)
    rect[1] = pts[np.argmin(diff)]  # Top-right
    rect[3] = pts[np.argmax(diff)]  # Bottom-left
    return rect
# improve image for ocr
def improve_image_for_ocr(image):
    """Cải thiện chất lượng ảnh trước khi OCR"""
    # Chuyển sang ảnh xám
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Tăng độ tương phản
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    contrast = clahe.apply(gray)
    
    # Khử nhiễu
    denoised = cv2.fastNlMeansDenoising(contrast)
    
    # Làm sắc nét
    kernel = np.array([[-1,-1,-1],
                      [-1, 9,-1],
                      [-1,-1,-1]])
    sharpened = cv2.filter2D(denoised, -1, kernel)
    
    # Áp dụng ngưỡng thích ứng
    thresh = cv2.adaptiveThreshold(
        sharpened,
        255,
        cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY,
        21,  # Tăng kích thước block
        11   # Tăng hằng số C
    )
    
    # Dilation để làm dày chữ
    kernel = np.ones((2,2), np.uint8)
    dilated = cv2.dilate(thresh, kernel, iterations=1)
    
    return dilated

def perform_google_ocr(image):
    try:
        credentials = service_account.Credentials.from_service_account_file(
            'keys/api-key.json')
        client = vision.ImageAnnotatorClient(credentials=credentials)
        success, encoded_image = cv2.imencode('.png', image)
        if not success:
            raise Exception("Failed to encode image")
            
        content = encoded_image.tobytes()
        image = vision.Image(content=content)

        try:
            response = client.text_detection(image=image)
            if response.error.message:
                raise Exception(
                    '{}\nFor more info on error messages, check: '
                    'https://cloud.google.com/apis/design/errors'.format(
                        response.error.message))
                        
            texts = response.text_annotations
            if not texts:
                print("No text detected")
                return [], ""
            
            # Lấy toàn bộ văn bản từ text_annotations[0]
            full_text = texts[0].description if texts else ""
            
            # Xử lý từng block text riêng biệt
            results = []
            for text in texts[1:]:  # Bỏ qua text đầu tiên vì đã lấy ở trên
                vertices = [(vertex.x, vertex.y) 
                    for vertex in text.bounding_poly.vertices]
                results.append((vertices, text.description))
                
            return results, full_text

        except Exception as e:
            print(f"API Error: {str(e)}")
            return [], ""
            
    except Exception as e:
        print(f"OCR Error: {str(e)}")
        return [], ""

def enhance_image(image):
    """Cải thiện chất lượng ảnh"""
    # Tăng độ tương phản
    lab = cv2.cvtColor(image, cv2.COLOR_BGR2LAB)
    l, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=3.0, tileGridSize=(8,8))
    cl = clahe.apply(l)
    enhanced = cv2.merge((cl,a,b))
    enhanced = cv2.cvtColor(enhanced, cv2.COLOR_LAB2BGR)
    
    # Giảm nhiễu
    enhanced = cv2.fastNlMeansDenoisingColored(enhanced, None, 10, 10, 7, 21)
    return enhanced

def process_scanned_image(image):
    """Xử lý ảnh scan sang dạng trắng đen"""
    # Chuyển sang ảnh xám
    gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    
    # Tăng độ tương phản
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
    contrast = clahe.apply(gray)
    
    # Áp dụng ngưỡng thích ứng để chuyển sang trắng đen
    binary = cv2.adaptiveThreshold(
        contrast,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
        21,
        11
    )
    
    # Chuyển lại thành ảnh BGR để lưu
    binary_bgr = cv2.cvtColor(binary, cv2.COLOR_GRAY2BGR)
    return binary_bgr

def process_image_async(image, filename):
    """Xử lý ảnh bất đồng bộ"""
    try:
        # Hiển thị thanh tiến trình
        progress = tk.Toplevel(giaodien)
        progress.title("Đang xử lý...")
        progress_label = tk.Label(progress, text="Đang xử lý ảnh...")
        progress_label.pack(pady=10)
        
        # Tạo bản sao cho xử lý scan
        scan_image = image.copy()
        
        # Cải thiện chất lượng ảnh cho scan
        enhanced = enhance_image(scan_image)
        #show_image(enhanced, "Enhanced Image", giaodien)
        
        # Tìm góc tài liệu và thực hiện scan
        corners = find_document_corners(enhanced)
        if corners is None:
            messagebox.showwarning("Cảnh báo", "Không tìm thấy biên tài liệu. Xử lý toàn bộ ảnh.")
            scanned_image = enhanced
        else:
            # Vẽ contour và thực hiện perspective transform
            cv2.drawContours(enhanced, [corners], -1, (0, 255, 0), 3)
            #show_image(enhanced, "Detected Document", giaodien)
            warped = perform_perspective_transform(enhanced, corners)
            
            # Xử lý ảnh scan sang dạng trắng đen
            scanned_image = process_scanned_image(warped)
            
        # Thực hiện OCR trên ảnh gốc
        progress_label.config(text="Đang nhận dạng văn bản...")
        results, full_text = perform_google_ocr(image)  # Sử dụng ảnh gốc cho OCR
        
        if results:
            # Xử lý và hiển thị kết quả
            output_image = process_ocr_results(image.copy(), results)  # Vẽ boxes trên ảnh gốc
            
            # Lưu cả ảnh scan và kết quả OCR
            save_all_results(full_text, scanned_image, output_image, filename)
            messagebox.showinfo("Thành công", "Đã hoàn thành quét và nhận dạng văn bản!")
        else:
            messagebox.showerror("Lỗi", "Không tìm thấy văn bản trong ảnh!")
            
        progress.destroy()
        
    except Exception as e:
        messagebox.showerror("Lỗi", f"Lỗi khi xử lý ảnh: {str(e)}")
        if 'progress' in locals():
            progress.destroy()

def save_all_results(text, scanned_image, ocr_image, filename):
    """Lưu tất cả kết quả bao gồm ảnh scan và OCR"""
    base_name = os.path.splitext(os.path.basename(filename))[0]
    output_dir = "output"
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
            
    # Lưu text
    text_path = os.path.join(output_dir, f"{base_name}_ocr.txt")
    with open(text_path, "w", encoding="utf-8") as f:
        f.write(text)
        
    # Lưu ảnh đã scan (trắng đen)
    scan_path = os.path.join(output_dir, f"{base_name}_scanned.jpg")
    cv2.imwrite(scan_path, scanned_image)
    
    # Lưu ảnh có boxes OCR
    ocr_path = os.path.join(output_dir, f"{base_name}_ocr_boxes.jpg")
    cv2.imwrite(ocr_path, ocr_image)

    # Hiển thị cửa sổ kết quả với cả hai ảnh
    show_results_window(text, scan_path, ocr_path)

def save_as_word(text, filename):
    """Lưu văn bản sang định dạng Word"""
    try:
        from docx import Document
        doc = Document()
        doc.add_paragraph(text)
        doc.save(filename)
        return True
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể lưu file Word: {str(e)}")
        return False

def show_results_window(text, scan_path, ocr_path):
    """Hiển thị cửa sổ kết quả với khả năng chỉnh sửa văn bản"""
    results_window = tk.Toplevel(giaodien)
    results_window.title("Kết Quả Quét Văn Bản")
    results_window.geometry("1024x768")
    
    # Tạo style cho các widget
    style = ttk.Style()
    style.configure("Title.TLabel", font=("Times New Roman", 16, "bold"))
    style.configure("Content.TFrame", padding=10)
    style.configure("Button.TButton", font=("Times New Roman", 11))

    main_frame = ttk.Frame(results_window, style="Content.TFrame")
    main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

    # Notebook với 3 tab riêng biệt
    notebook = ttk.Notebook(main_frame)
    notebook.pack(fill=tk.BOTH, expand=True, pady=(10, 20))

    # 1. Tab ảnh đã scan
    scan_frame = ttk.Frame(notebook, style="Content.TFrame")
    notebook.add(scan_frame, text=" Ảnh đã scan ")
    
    scan_label = ttk.Label(scan_frame, text="Kết quả scan tài liệu (Trắng đen)", style="Title.TLabel")
    scan_label.pack(pady=(5,10))
    
    # Canvas cho ảnh scan với thanh cuộn
    scan_canvas = tk.Canvas(scan_frame, bg='white')  # Thêm nền trắng
    scan_scrollbar_y = ttk.Scrollbar(scan_frame, orient=tk.VERTICAL, command=scan_canvas.yview)
    scan_scrollbar_x = ttk.Scrollbar(scan_frame, orient=tk.HORIZONTAL, command=scan_canvas.xview)
    
    scan_canvas.configure(yscrollcommand=scan_scrollbar_y.set, xscrollcommand=scan_scrollbar_x.set)
    
    scan_scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    scan_scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    scan_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    scan_container = ttk.Frame(scan_canvas)
    scan_canvas.create_window((0, 0), window=scan_container, anchor='nw')

    try:
        scan_img = Image.open(scan_path)
        # Giữ tỷ lệ ảnh khi resize
        display_width = 900
        ratio = display_width / scan_img.width
        display_height = int(scan_img.height * ratio)
        scan_img = scan_img.resize((display_width, display_height), Image.Resampling.LANCZOS)
        scan_photo = ImageTk.PhotoImage(scan_img)
        
        scan_img_label = ttk.Label(scan_container, image=scan_photo)
        scan_img_label.image = scan_photo
        scan_img_label.pack(pady=10)
        
        # Cập nhật scroll region
        scan_container.update_idletasks()
        scan_canvas.configure(scrollregion=scan_canvas.bbox('all'))
    except Exception as e:
        error_label = ttk.Label(
            scan_container,
            text=f"Không thể hiển thị ảnh scan: {str(e)}",
            style="Title.TLabel"
        )
        error_label.pack(pady=10)

    # 2. Tab ảnh OCR
    ocr_frame = ttk.Frame(notebook, style="Content.TFrame")
    notebook.add(ocr_frame, text=" Ảnh OCR ")
    
    ocr_label = ttk.Label(ocr_frame, text="Kết quả nhận dạng văn bản", style="Title.TLabel")
    ocr_label.pack(pady=(5,10))
    
    # Canvas cho ảnh OCR với thanh cuộn
    ocr_canvas = tk.Canvas(ocr_frame)
    ocr_scrollbar_y = ttk.Scrollbar(ocr_frame, orient=tk.VERTICAL, command=ocr_canvas.yview)
    ocr_scrollbar_x = ttk.Scrollbar(ocr_frame, orient=tk.HORIZONTAL, command=ocr_canvas.xview)
    
    ocr_canvas.configure(yscrollcommand=ocr_scrollbar_y.set, xscrollcommand=ocr_scrollbar_x.set)
    
    ocr_scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    ocr_scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    ocr_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    ocr_container = ttk.Frame(ocr_canvas)
    ocr_canvas.create_window((0, 0), window=ocr_container, anchor='nw')

    try:
        ocr_img = Image.open(ocr_path)
        # Giữ tỷ lệ ảnh
        display_width = 900
        ratio = display_width / ocr_img.width
        display_height = int(ocr_img.height * ratio)
        ocr_img = ocr_img.resize((display_width, display_height), Image.Resampling.LANCZOS)
        ocr_photo = ImageTk.PhotoImage(ocr_img)
        
        ocr_img_label = ttk.Label(ocr_container, image=ocr_photo)
        ocr_img_label.image = ocr_photo
        ocr_img_label.pack(pady=10)
        
        ocr_container.update_idletasks()
        ocr_canvas.configure(scrollregion=ocr_canvas.bbox('all'))
    except Exception as e:
        error_label = ttk.Label(
            ocr_container,
            text=f"Không thể hiển thị ảnh OCR: {str(e)}",
            style="Title.TLabel"
        )
        error_label.pack(pady=10)

    # 3. Tab văn bản
    text_frame = ttk.Frame(notebook, style="Content.TFrame")
    notebook.add(text_frame, text=" Văn bản ")
    
    text_area = tk.Text(
        text_frame,
        wrap=tk.WORD,
        font=("Times New Roman", 12),
        padx=10,
        pady=10,
        undo=True  # Cho phép undo/redo
    )

    # Thêm cả scrollbar dọc và ngang
    text_scrollbar_y = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_area.yview)
    text_scrollbar_x = ttk.Scrollbar(text_frame, orient=tk.HORIZONTAL, command=text_area.xview)

    # Cấu hình text area với cả hai scrollbar
    text_area.configure(
        yscrollcommand=text_scrollbar_y.set,
        xscrollcommand=text_scrollbar_x.set,
        wrap='none'  # Cho phép cuộn ngang
    )

    # Pack các thành phần
    text_scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    text_scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    text_area.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

    # Thêm thanh công cụ cho text editor
    editor_toolbar = ttk.Frame(text_frame)
    editor_toolbar.pack(fill=tk.X, before=text_area, pady=(0, 5))
    
    # Nút Undo/Redo
    ttk.Button(
        editor_toolbar,
        text="⟲ Hoàn tác",
        command=lambda: text_area.edit_undo() if text_area.edit_modified() else None
    ).pack(side=tk.LEFT, padx=2)
    
    ttk.Button(
        editor_toolbar,
        text="⟳ Làm lại",
        command=lambda: text_area.edit_redo()
    ).pack(side=tk.LEFT, padx=2)
    
    # Nút lưu thay đổi
    def save_changes():
        try:
            modified_text = text_area.get("1.0", tk.END)
            save_path = filedialog.asksaveasfilename(
                defaultextension=".txt",
                filetypes=[
                    ("Text files", "*.txt"),
                    ("Word files", "*.docx"),
                    ("All files", "*.*")
                ],
                title="Lưu văn bản đã chỉnh sửa"
            )
            if save_path:
                # Lưu theo định dạng file
                if save_path.endswith('.docx'):
                    save_as_word(modified_text, save_path)
                else:
                    with open(save_path, 'w', encoding='utf-8') as f:
                        f.write(modified_text)
                messagebox.showinfo("Thành công", "Đã lưu văn bản!")
                # Cập nhật trạng thái đã lưu
                text_area.edit_modified(False)
                
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể lưu văn bản: {str(e)}")

    save_button = ttk.Button(
        editor_toolbar,
        text="💾 Lưu thay đổi",
        command=save_changes
    )
    save_button.pack(side=tk.LEFT, padx=2)
    
    # Thêm nút định dạng cơ bản
    ttk.Button(
        editor_toolbar,
        text="In đậm",
        command=lambda: format_text(text_area, "bold")
    ).pack(side=tk.LEFT, padx=2)
    
    ttk.Button(
        editor_toolbar,
        text="In nghiêng",
        command=lambda: format_text(text_area, "italic")
    ).pack(side=tk.LEFT, padx=2)

    # Cảnh báo khi có thay đổi chưa lưu
    def on_text_change(event=None):
        if text_area.edit_modified():
            save_button.configure(style="Accent.TButton")
        
    text_area.bind("<<Modified>>", on_text_change)
    
    # Thêm văn bản và cấu hình scrollbar như cũ
    text_area.insert("1.0", text)
    text_area.configure(yscrollcommand=text_scrollbar_y.set, xscrollcommand=text_scrollbar_x.set)
    
    # Cập nhật nút trong text_buttons_frame
    text_buttons_frame = ttk.Frame(text_frame)
    text_buttons_frame.pack(pady=10)
    
    ttk.Button(
        text_buttons_frame,
        text="Sao chép văn bản",
        command=lambda: copy_to_clipboard(text_area.get("1.0", tk.END))
    ).pack(side=tk.LEFT, padx=5)

    # Frame chứa các nút điều khiển
    button_frame = ttk.Frame(main_frame)
    button_frame.pack(fill=tk.X, pady=(10, 0))

    # Frame cho các nút văn bản
    text_buttons_frame = ttk.LabelFrame(button_frame, text="Văn bản")
    text_buttons_frame.pack(side=tk.LEFT, padx=5, pady=5)

    # Thêm nút sao chép
    ttk.Button(
        text_buttons_frame,
        text="Sao chép văn bản",
        command=lambda: copy_to_clipboard(text),
        style="Button.TButton"
    ).pack(side=tk.LEFT, padx=2)

    # Thêm dropdown và nút lưu
    create_save_menu(text_buttons_frame, text)

    # Các nút cho ảnh
    image_buttons_frame = ttk.LabelFrame(button_frame, text="Ảnh")
    image_buttons_frame.pack(side=tk.LEFT, padx=5, pady=5)

    def save_scanned_image():
        """Lưu ảnh đã scan"""
        filetypes = [
            ("JPEG files", "*.jpg"),
            ("PNG files", "*.png"),
            ("All files", "*.*")
        ]
        filename = filedialog.asksaveasfilename(
            defaultextension=".jpg",
            filetypes=filetypes,
            initialfile="anh_da_scan"
        )
        if filename:
            try:
                import shutil
                shutil.copy2(scan_path, filename)
                messagebox.showinfo("Thành công", "Đã lưu ảnh đã scan!")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Lỗi khi lưu ảnh: {str(e)}")

    def save_ocr_image():
        """Lưu ảnh có boxes OCR"""
        filetypes = [
            ("JPEG files", "*.jpg"),
            ("PNG files", "*.png"),
            ("All files", "*.*")
        ]
        filename = filedialog.asksaveasfilename(
            defaultextension=".jpg",
            filetypes=filetypes,
            initialfile="anh_ocr"
        )
        if filename:
            try:
                import shutil
                shutil.copy2(ocr_path, filename)
                messagebox.showinfo("Thành công", "Đã lưu ảnh OCR!")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Lỗi khi lưu ảnh: {str(e)}")

    # Nút đóng
    ttk.Button(
        button_frame,
        text="Đóng",
        command=results_window.destroy,
        style="Button.TButton"
    ).pack(side=tk.RIGHT, padx=5)

def process_ocr_results(image, results):
    """Xử lý kết quả OCR"""
    # Chỉ vẽ các khung và text lên ảnh, không tạo output_text
    for (bbox, text) in results:
        tl = (int(bbox[0][0]), int(bbox[0][1]))
        br = (int(bbox[2][0]), int(bbox[2][1]))
        
        # Vẽ bbox và text
        cv2.rectangle(image, tl, br, (0, 255, 0), 2)
        cv2.putText(image, text, 
                   (tl[0], tl[1] - 10),
                   cv2.FONT_HERSHEY_SIMPLEX, 0.8, (0, 255, 0), 2)
    
    return image

def perform_perspective_transform(image, corners):
    """Cải thiện perspective transform"""
    # Sắp xếp các điểm góc
    rect = order_points(corners.reshape(4, 2))
    (tl, tr, br, bl) = rect
    
    # Tính toán kích thước đầu ra tối ưu
    width_a = np.sqrt(((br[0] - bl[0]) ** 2) + ((br[1] - bl[1]) ** 2))
    width_b = np.sqrt(((tr[0] - tl[0]) ** 2) + ((tr[1] - tl[1]) ** 2))
    max_width = max(int(width_a), int(width_b))
    
    height_a = np.sqrt(((tr[0] - br[0]) ** 2) + ((tr[1] - br[1]) ** 2))
    height_b = np.sqrt(((tl[0] - bl[0]) ** 2) + ((tl[1] - bl[1]) ** 2))
    max_height = max(int(height_a), int(height_b))
    
    # Đảm bảo kích thước đầu ra hợp lý
    if max_width > max_height:
        max_height = int(max_width * 1.414)  # Tỷ lệ A4
    else:
        max_width = int(max_height / 1.414)
    
    # Tạo ma trận biến đổi
    dst = np.array([
        [0, 0],
        [max_width - 1, 0],
        [max_width - 1, max_height - 1],
        [0, max_height - 1]
    ], dtype="float32")
    
    # Thực hiện biến đổi perspective
    M = cv2.getPerspectiveTransform(rect, dst)
    warped = cv2.warpPerspective(image, M, (max_width, max_height))
    
    return warped

class CropFrame:
    def __init__(self, master, image_path):
        self.top = tk.Toplevel(master)
        self.top.title("Điều chỉnh vùng quét")
        
        # Lấy kích thước màn hình
        screen_width = self.top.winfo_screenwidth()
        screen_height = self.top.winfo_screenheight()
        
        # Đọc ảnh gốc
        self.original = cv2.imread(image_path)
        
        # Tính toán tỷ lệ scale phù hợp với màn hình
        max_display_width = int(screen_width * 0.8)  # 80% chiều rộng màn hình
        max_display_height = int(screen_height * 0.8)  # 80% chiều cao màn hình
        
        # Tính toán kích thước window
        img_h, img_w = self.original.shape[:2]
        width_ratio = max_display_width / img_w
        height_ratio = max_display_height / img_h
        self.scale_factor = min(width_ratio, height_ratio)
        
        # Resize ảnh theo tỷ lệ
        new_width = int(img_w * self.scale_factor)
        new_height = int(img_h * self.scale_factor)
        self.image = cv2.resize(self.original, (new_width, new_height))
        self.display_image = self.image.copy()
        
        # Cấu hình window
        window_width = new_width + 40  # Thêm padding
        window_height = new_height + 100  # Thêm space cho buttons
        
        # Căn giữa cửa sổ
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2
        self.top.geometry(f"{window_width}x{window_height}+{x}+{y}")
        
        # Cho phép resize window
        self.top.resizable(True, True)
        
        # Tạo main frame với grid layout
        main_frame = ttk.Frame(self.top)
        main_frame.grid(row=0, column=0, sticky="nsew")
        
        # Cấu hình grid
        self.top.grid_rowconfigure(0, weight=1)
        self.top.grid_columnconfigure(0, weight=1)
        
        # Canvas với scrollbars
        self.canvas = tk.Canvas(main_frame)
        h_scrollbar = ttk.Scrollbar(main_frame, orient=tk.HORIZONTAL, command=self.canvas.xview)
        v_scrollbar = ttk.Scrollbar(main_frame, orient=tk.VERTICAL, command=self.canvas.yview)
        
        # Cấu hình canvas
        self.canvas.configure(xscrollcommand=h_scrollbar.set, yscrollcommand=v_scrollbar.set)
        
        # Grid layout cho scrollbars và canvas
        self.canvas.grid(row=0, column=0, sticky="nsew")
        h_scrollbar.grid(row=1, column=0, sticky="ew")
        v_scrollbar.grid(row=0, column=1, sticky="ns")
        
        # Cấu hình grid weights
        main_frame.grid_rowconfigure(0, weight=1)
        main_frame.grid_columnconfigure(0, weight=1)
        
        # Frame cho buttons
        button_frame = ttk.Frame(self.top)
        button_frame.grid(row=1, column=0, pady=5, sticky="ew")
        
        # Style cho buttons
        style = ttk.Style()
        style.configure("Action.TButton", padding=5)
        
        # Buttons
        ttk.Button(
            button_frame,
            text="Tự động phát hiện",
            command=self.auto_detect,
            style="Action.TButton"
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            button_frame,
            text="Quét vùng đã chọn",
            command=self.confirm_crop,
            style="Action.TButton"
        ).pack(side=tk.LEFT, padx=5)
        
        ttk.Button(
            button_frame,
            text="Hủy",
            command=self.top.destroy,
            style="Action.TButton"
        ).pack(side=tk.RIGHT, padx=5)
        
        # Bind events
        self.canvas.bind('<Configure>', self.on_resize)
        self.canvas.bind('<Button-1>', self.on_click)
        self.canvas.bind('<B1-Motion>', self.on_drag)
        self.canvas.bind('<ButtonRelease-1>', self.on_release)
        
        # Các biến cho việc crop
        self.points = []
        self.dragging = None
        self.drag_threshold = 20
        
        # Hiển thị ảnh ban đầu
        self.update_display()
        self.auto_detect()

    def on_resize(self, event):
        """Xử lý khi window được resize"""
        # Cập nhật scroll region
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        
        # Tính toán lại kích thước hiển thị
        canvas_width = event.width
        canvas_height = event.height
        
        # Resize ảnh hiển thị nếu cần
        if hasattr(self, 'display_image'):
            self.update_display()

    def update_display(self):
        """Cập nhật hiển thị ảnh với responsive"""
        self.display_image = self.image.copy()
        
        if len(self.points) > 0:
            # Scale các điểm theo tỷ lệ hiện tại
            scaled_points = [(int(p[0]), int(p[1])) for p in self.points]
            
            # Vẽ các điểm và đường nối
            for pt in scaled_points:
                cv2.circle(self.display_image, pt, 5, (0, 255, 0), -1)
            
            pts = np.array(scaled_points, np.int32)
            cv2.polylines(self.display_image, [pts], True, (0, 255, 0), 2)
        
        # Chuyển đổi và hiển thị
        photo = convert_cv_to_tkinter(self.display_image)
        self.canvas.delete("all")
        self.canvas.create_image(0, 0, image=photo, anchor='nw')
        self.canvas.image = photo
        
        # Cập nhật scroll region
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))

    def find_closest_point(self, x, y):
        """Tìm điểm gần nhất với tọa độ (x,y)"""
        if not self.points:
            return None
        distances = [np.sqrt((pt[0]-x)**2 + (pt[1]-y)**2) for pt in self.points]
        min_dist = min(distances)
        if min_dist < self.drag_threshold:
            return distances.index(min_dist)
        return None

    def on_click(self, event):
        """Xử lý sự kiện click chuột"""
        self.dragging = self.find_closest_point(event.x, event.y)

    def on_drag(self, event):
        """Xử lý sự kiện kéo chuột"""
        if self.dragging is not None:
            self.points[self.dragging] = [event.x, event.y]
            self.update_display()

    def on_release(self, event):
        """Xử lý sự kiện thả chuột"""
        self.dragging = None

    def confirm_crop(self):
        """Xác nhận và thực hiện crop ảnh"""
        if len(self.points) == 4:
            # Chuyển đổi tỷ lệ về ảnh gốc
            h_orig, w_orig = self.original.shape[:2]
            h_disp, w_disp = self.image.shape[:2]
            
            scale_x = w_orig / w_disp
            scale_y = h_orig / h_disp
            
            points_orig = np.array([[int(p[0] * scale_x), int(p[1] * scale_y)] 
                                  for p in self.points], dtype=np.float32)
            
            # Thực hiện transform
            warped = perform_perspective_transform(self.original, points_orig)
            
            # Lưu kết quả và đóng cửa sổ crop
            self.result = warped
            self.top.destroy()

    def auto_detect(self):
        """Tự động phát hiện góc tài liệu"""
        try:
            # Tạo bản sao ảnh để xử lý
            img_to_detect = self.image.copy()
            
            # Phát hiện góc
            corners = find_document_corners(img_to_detect)
            
            if corners is not None:
                # Chuyển đổi corners thành list các điểm
                self.points = corners.reshape(4, 2).tolist()
                self.update_display()
            else:
                # Nếu không phát hiện được, dùng góc mặc định
                h, w = self.image.shape[:2]
                margin = min(w, h) // 10  # Margin 10% 
                self.points = [
                    [margin, margin],  # Top-left
                    [w - margin, margin],  # Top-right
                    [w - margin, h - margin],  # Bottom-right
                    [margin, h - margin]  # Bottom-left
                ]
                self.update_display()
                messagebox.showinfo(
                    "Thông báo",
                    "Không thể tự động phát hiện góc, vui lòng điều chỉnh thủ công"
                )
        except Exception as e:
            messagebox.showerror(
                "Lỗi",
                f"Lỗi khi tự động phát hiện góc: {str(e)}\nVui lòng điều chỉnh thủ công"
            )
            # Khởi tạo điểm mặc định
            h, w = self.image.shape[:2]
            margin = min(w, h) // 10
            self.points = [
                [margin, margin],
                [w - margin, margin], 
                [w - margin, h - margin],
                [margin, h - margin]
            ]
            self.update_display()

def select_and_scan():
    """Hàm xử lý chọn và quét ảnh với crop"""
    filetypes = (
        ('Image files', '*.jpg *.jpeg *.png *.bmp'),
        ('All files', '*.*')
    )
    
    filename = filedialog.askopenfilename(
        title='Chọn ảnh cần quét',
        initialdir='./images',
        filetypes=filetypes
    )
    
    if filename:
        try:
            # Hiển thị cửa sổ crop
            crop_window = CropFrame(giaodien, filename)
            giaodien.wait_window(crop_window.top)
            
            # Nếu đã crop thành công
            if hasattr(crop_window, 'result'):
                # Xử lý ảnh đã crop trong thread riêng
                thread = threading.Thread(
                    target=process_image_async,
                    args=(crop_window.result, filename)
                )
                thread.start()
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Lỗi: {str(e)}")
    else:
        messagebox.showinfo("Thông báo", "Không có file nào được chọn")

def save_as_pdf(text, filename):
    """Lưu văn bản sang định dạng PDF"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        
        # Đăng ký font Unicode để hỗ trợ tiếng Việt
        try:
            pdfmetrics.registerFont(TTFont('Times', 'times.ttf'))
        except:
            # Nếu không có font Times, dùng font mặc định
            font_name = 'Helvetica'
        else:
            font_name = 'Times'
        
        # Tạo file PDF
        c = canvas.Canvas(filename, pagesize=A4)
        width, height = A4
        
        # Cấu hình font và cỡ chữ
        c.setFont(font_name, 12)
        
        # Chia văn bản thành các dòng để vẽ
        y = height - 50  # Để lại lề trên
        margin = 50  # Lề trái và phải
        line_height = 20
        
        # Tách văn bản thành các dòng
        words = text.split()
        current_line = []
        for word in words:
            current_line.append(word)
            # Kiểm tra độ rộng của dòng hiện tại
            line_width = c.stringWidth(' '.join(current_line), font_name, 12)
            if line_width > width - 2*margin:
                # Nếu dòng quá dài, vẽ dòng trước đó
                current_line.pop()  # Bỏ từ cuối ra
                line = ' '.join(current_line)
                c.drawString(margin, y, line)
                y -= line_height
                # Bắt đầu dòng mới với từ bị tràn
                current_line = [word]
                
                # Kiểm tra nếu hết trang
                if y < margin:
                    c.showPage()
                    y = height - 50
                    c.setFont(font_name, 12)
        
        # Vẽ dòng cuối cùng
        if current_line:
            line = ' '.join(current_line)
            c.drawString(margin, y, line)
        
        c.save()
        return True
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể lưu file PDF: {str(e)}")
        return False

def create_save_menu(text_buttons_frame, text):
    """Tạo dropdown menu cho việc lưu văn bản"""
    # Tạo frame chứa combobox và nút lưu
    save_frame = ttk.Frame(text_buttons_frame)
    save_frame.pack(side=tk.LEFT, padx=5)

    # Tạo combobox để chọn định dạng
    format_var = tk.StringVar()
    format_combo = ttk.Combobox(
        save_frame, 
        textvariable=format_var,
        state='readonly',
        width=15
    )
    format_combo['values'] = [
        'Văn bản (.txt)',
        'PDF (.pdf)',
        'Word (.docx)',
        'Rich Text (.rtf)'
    ]
    format_combo.set('Văn bản (.txt)')  # Giá trị mặc định
    format_combo.pack(side=tk.LEFT, padx=2)

    def save_with_format():
        selected = format_combo.get()
        if selected == 'Văn bản (.txt)':
            extension = '.txt'
            file_type = 'Text files'
        elif selected == 'PDF (.pdf)':
            extension = '.pdf'
            file_type = 'PDF files'
        elif selected == 'Word (.docx)':
            extension = '.docx'
            file_type = 'Word files'
        else:  # RTF
            extension = '.rtf'
            file_type = 'Rich Text files'

        filename = filedialog.asksaveasfilename(
            defaultextension=extension,
            filetypes=[(file_type, f'*{extension}'), ('All files', '*.*')],
            initialfile=f"ket_qua_scan{extension}"
        )

        if filename:
            try:
                if extension == '.txt':
                    with open(filename, 'w', encoding='utf-8') as f:
                        f.write(text)
                elif extension == '.pdf':
                    if save_as_pdf(text, filename):
                        messagebox.showinfo("Thành công", "Đã lưu văn bản sang PDF!")
                        return
                elif extension == '.docx':
                    if save_as_word(text, filename):
                        messagebox.showinfo("Thành công", "Đã lưu văn bản sang Word!")
                        return
                elif extension == '.rtf':
                    if save_as_rtf(text, filename):
                        messagebox.showinfo("Thành công", "Đã lưu văn bản sang RTF!")
                        return
                messagebox.showinfo("Thành công", "Đã lưu văn bản!")
            except Exception as e:
                messagebox.showerror("Lỗi", f"Lỗi khi lưu văn bản: {str(e)}")

    # Nút lưu
    ttk.Button(
        save_frame,
        text="Lưu văn bản",
        command=save_with_format,
        style="Button.TButton"
    ).pack(side=tk.LEFT, padx=2)

def copy_to_clipboard(text):
    """Sao chép văn bản vào clipboard"""
    giaodien.clipboard_clear()
    giaodien.clipboard_append(text)
    messagebox.showinfo("Thành công", "Đã sao chép văn bản vào clipboard!")

def save_as_rtf(text, filename):
    """Lưu văn bản sang định dạng RTF"""
    try:
        from pyrtf import Document, Section, Paragraph, TEXT
        doc = Document()
        section = Section()
        doc.Sections.append(section)
        
        for line in text.split('\n'):
            p = Paragraph()
            p.append(TEXT(line))
            section.append(p)
            
        with open(filename, 'wb') as f:
            doc.write(f)
        return True
    except Exception as e:
        messagebox.showerror("Lỗi", f"Không thể lưu file RTF: {str(e)}")
        return False

def show_results(scan_path, text):
    """Hiển thị kết quả scan với khả năng zoom"""
    results_window = tk.Toplevel(giaodien)
    results_window.title("Kết quả quét văn bản")
    results_window.geometry("1000x800")

    # Tạo style cho các widget
    style = ttk.Style()
    style.configure("Title.TLabel", font=("Times New Roman", 16, "bold"))
    style.configure("Content.TFrame", padding=10)
    style.configure("Button.TButton", font=("Times New Roman", 11))

    main_frame = ttk.Frame(results_window, style="Content.TFrame")
    main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)

    # Notebook với các tab
    notebook = ttk.Notebook(main_frame)
    notebook.pack(fill=tk.BOTH, expand=True, pady=(10, 20))

    # Biến để lưu trữ tỷ lệ zoom
    zoom_scale = tk.DoubleVar(value=1.0)
    
    # 1. Tab ảnh đã scan với zoom
    scan_frame = ttk.Frame(notebook, style="Content.TFrame")
    notebook.add(scan_frame, text=" Ảnh đã scan ")
    
    # Frame cho controls zoom
    zoom_frame = ttk.Frame(scan_frame)
    zoom_frame.pack(fill=tk.X, pady=5)
    
    ttk.Label(zoom_frame, text="Zoom:").pack(side=tk.LEFT, padx=5)
    zoom_slider = ttk.Scale(
        zoom_frame,
        from_=0.5,
        to=3.0,
        orient=tk.HORIZONTAL,
        variable=zoom_scale,
        length=200
    )
    zoom_slider.pack(side=tk.LEFT, padx=5)
    
    # Canvas cho ảnh với scroll
    scan_canvas = tk.Canvas(scan_frame, bg='white')
    scan_scrollbar_y = ttk.Scrollbar(scan_frame, orient=tk.VERTICAL, command=scan_canvas.yview)
    scan_scrollbar_x = ttk.Scrollbar(scan_frame, orient=tk.HORIZONTAL, command=scan_canvas.xview)
    
    scan_canvas.configure(yscrollcommand=scan_scrollbar_y.set, xscrollcommand=scan_scrollbar_x.set)
    
    scan_scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    scan_scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    scan_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    # Frame cho nội dung
    scan_container = ttk.Frame(scan_canvas)
    scan_canvas.create_window((0, 0), window=scan_container, anchor='nw')

    # 2. Tab văn bản với zoom
    text_frame = ttk.Frame(notebook, style="Content.TFrame")
    notebook.add(text_frame, text=" Văn bản ")
    
    # Frame cho text controls
    text_control_frame = ttk.Frame(text_frame)
    text_control_frame.pack(fill=tk.X, pady=5)
    
    # Nút điều chỉnh cỡ chữ
    ttk.Label(text_control_frame, text="Cỡ chữ:").pack(side=tk.LEFT, padx=5)
    font_size = tk.IntVar(value=12)
    
    def change_font_size(event=None):
        size = font_size.get()
        text_widget.configure(font=("Times New Roman", size))
    
    font_scale = ttk.Scale(
        text_control_frame,
        from_=8,
        to=32,
        variable=font_size,
        orient=tk.HORIZONTAL,
        length=200,
        command=change_font_size
    )
    font_scale.pack(side=tk.LEFT, padx=5)
    
    # Text widget với scrollbar
    text_container = ttk.Frame(text_frame)
    text_container.pack(fill=tk.BOTH, expand=True)
    
    text_widget = tk.Text(
        text_container,
        wrap=tk.WORD,
        font=("Times New Roman", 12),
        padx=10,
        pady=10,
        undo=True  # Cho phép undo/redo
    )
    text_scrollbar_y = ttk.Scrollbar(text_container, orient=tk.VERTICAL, command=text_widget.yview)
    text_scrollbar_x = ttk.Scrollbar(text_frame, orient=tk.HORIZONTAL, command=text_widget.xview)
    
    text_widget.configure(yscrollcommand=text_scrollbar_y.set, xscrollcommand=text_scrollbar_x.set)
    
    text_scrollbar_y.pack(side=tk.RIGHT, fill=tk.Y)
    text_scrollbar_x.pack(side=tk.BOTTOM, fill=tk.X)
    text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    
    # Thêm văn bản vào widget
    text_widget.insert('1.0', text)
    
    # Bind phím tắt Ctrl + Mouse wheel cho zoom
    def on_mousewheel(event):
        if event.state == 4:  # Ctrl
            if event.delta > 0:
                current = font_size.get()
                font_size.set(min(current + 1, 32))
            else:
                current = font_size.get()
                font_size.set(max(current - 1, 8))
            change_font_size()
            return "break"
    
    text_widget.bind('<Control-MouseWheel>', on_mousewheel)
    
    # Nút copy và lưu
    text_buttons_frame = ttk.Frame(text_frame)
    text_buttons_frame.pack(pady=10)
    
    ttk.Button(
        text_buttons_frame,
        text="Sao chép văn bản",
        command=lambda: copy_to_clipboard(text),
        style="Button.TButton"
    ).pack(side=tk.LEFT, padx=5)
    
    # Menu lưu văn bản
    create_save_menu(text_buttons_frame, text)
    
    # Cập nhật hiển thị ảnh scan với zoom
    def update_scan_image(event=None):
        try:
            scan_img = Image.open(scan_path)
            # Tính toán kích thước hiển thị với zoom
            base_width = 900
            zoom = zoom_scale.get()
            display_width = int(base_width * zoom)
            ratio = display_width / scan_img.width
            display_height = int(scan_img.height * ratio)
            
            scan_img = scan_img.resize((display_width, display_height), Image.Resampling.LANCZOS)
            scan_photo = ImageTk.PhotoImage(scan_img)
            
            # Xóa ảnh cũ nếu có
            for widget in scan_container.winfo_children():
                widget.destroy()
            
            scan_img_label = ttk.Label(scan_container, image=scan_photo)
            scan_img_label.image = scan_photo
            scan_img_label.pack(pady=10)
            
            # Cập nhật scroll region
            scan_canvas.configure(scrollregion=scan_canvas.bbox("all"))
            
        except Exception as e:
            messagebox.showerror("Lỗi", f"Không thể hiển thị ảnh: {str(e)}")
    
    # Bind sự kiện zoom cho ảnh
    zoom_slider.configure(command=update_scan_image)
    update_scan_image()  # Hiển thị ảnh ban đầu

def format_text(text_widget, format_type):
    """Hàm hỗ trợ định dạng văn bản cơ bản"""
    try:
        selected = text_widget.tag_ranges(tk.SEL)
        if selected:
            current_tags = text_widget.tag_names(tk.SEL_FIRST)
            
            if format_type == "bold":
                if "bold" in current_tags:
                    text_widget.tag_remove("bold", tk.SEL_FIRST, tk.SEL_LAST)
                else:
                    text_widget.tag_add("bold", tk.SEL_FIRST, tk.SEL_LAST)
                    text_widget.tag_configure("bold", font=("Times New Roman", 12, "bold"))
                    
            elif format_type == "italic":
                if "italic" in current_tags:
                    text_widget.tag_remove("italic", tk.SEL_FIRST, tk.SEL_LAST)
                else:
                    text_widget.tag_add("italic", tk.SEL_FIRST, tk.SEL_LAST)
                    text_widget.tag_configure("italic", font=("Times New Roman", 12, "italic"))
                    
    except tk.TclError:
        pass  # Không có text được chọn

def create_main_window():
    global giaodien, background_image
    
    # Tạo cửa sổ chính
    giaodien = tk.Tk()
    giaodien.title("Document Scanner")
    giaodien.geometry("1322x743")

    try:
        # Thử load ảnh nền
        bia = Image.open("images/wall.jpg")
        resize_image = bia.resize((1322, 743))
        background_image = ImageTk.PhotoImage(resize_image)
        
        # Tạo label với ảnh nền
        img_label = Label(giaodien, image=background_image)
        img_label.grid(column=0, row=0)
        
    except FileNotFoundError:
        # Nếu không tìm thấy file ảnh, tạo background màu
        giaodien.configure(bg='#f0f0f0')  # Màu xám nhạt
        
        # Tạo frame chứa nội dung
        main_frame = ttk.Frame(giaodien, padding="20")
        main_frame.grid(column=0, row=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        
        # Thêm tiêu đề
        title_label = ttk.Label(
            main_frame,
            text="Document Scanner",
            font=("Times New Roman", 24, "bold"),
            padding="10"
        )
        title_label.grid(column=0, row=0, pady=20)

    # Thêm nút scan
    scan_button = ttk.Button(
        giaodien if 'main_frame' not in locals() else main_frame,
        text="Chọn ảnh để quét và scan",
        style="Large.TButton",
        command=select_and_scan
    )
    scan_button.grid(column=0, row=1 if 'main_frame' in locals() else 0, pady=20)
    
    # Tạo style cho nút lớn
    style = ttk.Style()
    style.configure(
        "Large.TButton",
        font=("Times New Roman", 20),
        padding=10
    )

    return giaodien

# Thay thế đoạn code cũ bằng lời gọi hàm
giaodien = create_main_window()
giaodien.mainloop()